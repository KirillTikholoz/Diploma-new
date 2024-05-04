from docx import Document
import textract


# Функция для извлечения текста из документа .docx
# Функция для извлечения текста из документа .docx, включая таблицы
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + '\t'  # Добавляем табуляцию между ячейками таблицы
            text += '\n'
    return text.strip()


# Функция для извлечения текста из документа .doc
def extract_text_from_doc(file_path):
    text = textract.process(file_path).decode('utf-8')
    return text.strip()


print(extract_text_from_doc(
    "C:/Users/admin/Desktop/diploma/docs/%D0%BF%D1%80%D0%B8%D0%BB%D0%BE%D0%B6%D0%B5%D0%BD%D0%B8%D0%B5%20%D0%BA%20%D1%80%D0%B5%D1%88%D0%B5%D0%BD%D0%B8%D1%8E%2029-223.doc"))
