import os
import textract
from elasticsearch import Elasticsearch
from docx import Document
from docx import Document
import win32com.client


# Подключаемся к Elasticsearch
# es = Elasticsearch([{'host': 'localhost', 'port': 9200}])

# Функция для извлечения текста из документа .docx, включая таблицы
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + '\t'  # Добавляем табуляцию между ячейками таблицы
            text += '\n'
    return text.strip()


# Функция для извлечения текста из документа .doc
def extract_text_from_doc(file_path):
    text = textract.process(file_path).decode('utf-8')
    return text.strip()


# Функция для индексации документа в Elasticsearch
# def index_document(index_name, doc_id, text):
#     es.index(index=index_name, id=doc_id, body={'text': text})
#
# # Основная функция для индексации всех документов в заданной папке
# def index_documents_from_folder(folder_path, index_name):
#     for filename in os.listdir(folder_path):
#         if filename.endswith(".docx"):
#             file_path = os.path.join(folder_path, filename)
#             text = extract_text_from_docx(file_path)
#             index_document(index_name, filename, text)
#         elif filename.endswith(".doc"):
#             file_path = os.path.join(folder_path, filename)
#             text = extract_text_from_doc(file_path)
#             index_document(index_name, filename, text)

# Пример использования: индексация документов из папки 'docs' в индекс 'my_index'
# index_documents_from_folder('docs', 'my_index')
folder_path = r'C:\Users\admin\Desktop\diploma\docs'
for filename in os.listdir(folder_path):
    file_path = os.path.join(folder_path, filename)
    print(file_path)
    text = extract_text_from_docx(file_path)
    print(text)


# def index_documents_from_folder(folder_path, index):
#     for filename in os.listdir(folder_path):
#         if filename.endswith(".docx"):
#             file_path = os.path.join(folder_path, filename)
#             try:
#                 text = extract_text_from_docx(file_path)
#                 print(text)
#                 try:
#                     index_document(index, filename, text)
#                     print('Документ успешно проиндексирован:', filename)
#                 except Exception as e:
#                     print('Ошибка при индексации документа:', e)
#             except Exception as e:
#                 print('Ошибка при попытке извлечь текст:', e)
#
#         elif filename.endswith(".doc"):
#             try:
#                 file_path = os.path.join(folder_path, filename)
#                 text = extract_text_from_doc(file_path)
#                 index_document(index, filename, text)
#             except Exception as e:
#                 print('Формат doc:', e)

# def convert_doc_to_docx_with_pywin32(doc_file_path, docx_file_path):
#     word = win32com.client.Dispatch("Word.Application")
#     doc = word.Documents.Open(doc_file_path)
#     doc.SaveAs(docx_file_path, FileFormat=16)  # 16 - формат .docx
#     doc.Close()
#     word.Quit()
#
#
# def convert_doc_to_docx_2(doc_file_path, docx_file_path):
#     # Открываем документ .doc
#     doc = Document(doc_file_path)
#     # Создаем новый документ .docx
#     docx = Document()
#     # Копируем содержимое из документа .doc в новый документ .docx
#     for paragraph in doc.paragraphs:
#         docx.add_paragraph(paragraph.text)
#     # Сохраняем новый документ .docx
#     docx.save(docx_file_path)

# Основная функция для обработки всех файлов .doc в папке и их преобразования в .docx
# def batch_convert_docs_to_docx(input_folder, output_folder):
#     if not os.path.exists(output_folder):
#         os.makedirs(output_folder)
#
#     for filename in os.listdir(input_folder):
#         print(filename)
#         if filename.endswith(".doc"):
#             doc_file_path = os.path.join(input_folder, filename)
#             print(doc_file_path)
#             docx_file_path = os.path.join(output_folder, filename.replace(".doc", ".docx"))
#             print(docx_file_path)
#             convert_doc_to_docx(doc_file_path, docx_file_path)
#             print(f"Converted {filename} to .docx")