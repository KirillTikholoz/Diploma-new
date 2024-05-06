import os
import comtypes.client as cc
from docx import Document
import win32com.client


def convert_doc_to_docx_with_pywin32(doc_file_path, docx_file_path):
    # Создаем объект Word
    word = win32com.client.Dispatch("Word.Application")

    # Открываем файл .doc
    doc = word.Documents.Open(doc_file_path)

    # Конвертируем и сохраняем в формат .docx
    doc.SaveAs(docx_file_path, FileFormat=16)  # 16 - формат .docx

    # Закрываем файл .doc
    doc.Close()

    # Закрываем Word
    word.Quit()


def convert_doc_to_docx_2(doc_file_path, docx_file_path):
    # Открываем документ .doc
    doc = Document(doc_file_path)

    # Создаем новый документ .docx
    docx = Document()

    # Копируем содержимое из документа .doc в новый документ .docx
    for paragraph in doc.paragraphs:
        docx.add_paragraph(paragraph.text)

    # Сохраняем новый документ .docx
    docx.save(docx_file_path)


# Функция для преобразования файла .doc в .docx
def convert_doc_to_docx(doc_file_path, docx_file_path):
    word = cc.CreateObject("Word.Application")
    doc = word.Documents.Open(doc_file_path)
    print('файл открыт')
    doc.SaveAs(docx_file_path, FileFormat=16)  # 16 - формат .docx
    doc.Close()
    word.Quit()


# Основная функция для обработки всех файлов .doc в папке и их преобразования в .docx
def batch_convert_docs_to_docx(input_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(input_folder):
        print(filename)
        if filename.endswith(".doc"):
            doc_file_path = os.path.join(input_folder, filename)
            print(doc_file_path)
            docx_file_path = os.path.join(output_folder, filename.replace(".doc", ".docx"))
            print(docx_file_path)
            convert_doc_to_docx(doc_file_path, docx_file_path)
            print(f"Converted {filename} to .docx")
