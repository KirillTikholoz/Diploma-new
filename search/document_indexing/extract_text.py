from docx import Document
import textract


# Функция для извлечения текста из документа .docx, включая таблицы
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + '\t'  # Добавляем табуляцию между ячейками таблицы
            text += '\n'
    return text.strip()


# Функция для извлечения текста из документа .doc
def extract_text_from_doc(file_path):
    text = textract.process(file_path).decode('utf-8')
    return text.strip()
