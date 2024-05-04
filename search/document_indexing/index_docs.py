import os
import time
import textract
import pythoncom
# import threading

from elasticsearch import Elasticsearch
from docx import Document
from doc_to_docx import convert_doc_to_docx
from flask import Flask, request

app = Flask(__name__)

# Подключаемся к Elasticsearch
es = Elasticsearch("http://localhost:9200", request_timeout=60)

# Настройки индекса
index_settings = {
    "settings": {
        "analysis": {
            "analyzer": {
                "default": {
                    "type": "standard",  # Используем стандартный токенизатор
                    "lowercase": True  # Приводим все буквы к нижнему регистру
                }
            }
        }
    }
}


# Функция для извлечения текста из документа .docx, включая таблицы
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + '\t'  # Добавляем табуляцию между ячейками таблицы
            text += '\n'
    return text.strip()


# Функция для извлечения текста из документа .doc
def extract_text_from_doc(file_path):
    text = textract.process(file_path).decode('utf-8')
    return text.strip()


# Функция для индексации документа в Elasticsearch
def index_document(index_name, doc_id, text):
    es.index(index=index_name, id=doc_id, body={'text': text})


def extract_name_from_path(path):
    last_slash_position = max(path.rfind('/'), path.rfind('\\'))
    if last_slash_position != -1:
        # print(path.removeprefix(path[:last_slash_position+1]))
        return path.removeprefix(path[:last_slash_position+1])
    else:
        return path


def index_single_document(document_path, index_name='doc_index'):
    filename = extract_name_from_path(document_path)
    print(filename)
    if document_path.endswith(".docx"):
        try:
            text = extract_text_from_docx(document_path)
            print(text)
            try:
                index_document(index_name, filename, text)
                print('Документ успешно проиндексирован:', filename)
                time.sleep(1)
            except Exception as e:
                print('Ошибка при индексации документа:', e)
                time.sleep(1)
        except Exception as e:
            print('Ошибка при попытке извлечь текст:', e)
            time.sleep(1)
    elif filename.endswith(".doc"):
        print('Документ находится в старом формате и будет переведен в новый')
        docx_path = document_path[:-3] + 'docx'
        print(docx_path)
        try:
            # thread = threading.Thread(target=convert_doc_to_docx, args=(document_path, docx_path))
            pythoncom.CoInitialize()
            convert_doc_to_docx(document_path, docx_path)
            os.remove(document_path)
            index_single_document(docx_path, index_name)
        except Exception as e:
            print('Ошибка при переводе документа в новый формат:', e)


@app.route('/upload', methods=['POST'])
def upload_file():
    # print(request.form, '\n', request.files)
    # return 'кайф'
    # Получаем данные из формы
    file = request.files['file']
    id = request.form['id']

    # Обработка файла и id
    # Например, сохранение файла на сервере
    filepath = f'C:\\Users\\admin\\Desktop\\diploma\\docs\\{file.filename}'
    print(filepath)
    file.save(filepath)
    index_single_document(filepath)

    # Возвращаем сообщение об успешной загрузке
    return f'Файл успешно загружен и добавлен в поисковик по адресу {filepath}. ID:'.format(id)


def index_documents_from_folder(folder_path, index_name='doc_index'):
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            try:
                text = extract_text_from_docx(file_path)
                print(text)
                try:
                    index_document(index_name, filename, text)
                    print('Документ успешно проиндексирован:', filename)
                    time.sleep(1)
                except Exception as e:
                    print('Ошибка при индексации документа:', e)
                    time.sleep(1)
            except Exception as e:
                print('Ошибка при попытке извлечь текст:', e)
                time.sleep(1)

        elif filename.endswith(".doc"):
            try:
                file_path = os.path.join(folder_path, filename)
                text = extract_text_from_doc(file_path)
                index_document(index_name, filename, text)
            except Exception as e:
                print('Формат doc:', e)
                time.sleep(1)


if __name__ == "__main__":
    # index = 'doc_index'
    # Пример использования: индексация документов из папки 'docs' в индекс 'my_index'
    # es.indices.create(index=index, body=index_settings, ignore=400)
    # index_documents_from_folder(folder_path=r"C:\Users\admin\Desktop\diploma\docs", index_name=index)
    app.run(debug=True)
